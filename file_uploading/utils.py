import re
import csv
import datefinder
import pandas as pd
from docx import Document
import os
from distutils.log import debug 
from fileinput import filename 
from flask import *  

app = Flask(__name__)   

def clean_markdown_document(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()

    # Find the position of the first occurrence of "1."
    first_occurrence_index = content.find("1.")
    # Find the position of the first occurrence of a footnote
    first_occurrence_of_footnote = content.find("[^1]:")  

    # Determine the index to cut off the content
    if first_occurrence_index!= -1:
        if first_occurrence_of_footnote!= -1 and first_occurrence_of_footnote > first_occurrence_index:
            cleaned_content = content[:first_occurrence_of_footnote]  # Slice content up to the footnote
        else:
            cleaned_content = content[first_occurrence_index:]  # Keep content from "1." onwards
    else:
        cleaned_content = content  # If "1." is not found, keep the content as is

    with open(output_file, 'w', encoding='utf-8') as file:
        file.write(cleaned_content)

    print(f"Markdown document cleaned and saved as {output_file}")


def extract_numbered_items_to_csv(input_file, output_file):
    with open(input_file, 'r', encoding='utf-8') as file:
        content = file.read()

    # Regex pattern to find numbered items
    pattern = r'(\d+)\.\s*(.*?)\n(?=\d+\.\s|$)'  # Matches numbers followed by text, ending with a new line before another number

    # Find all matches
    matches = re.findall(pattern, content, re.DOTALL)

    # Write the matches to a CSV file
    with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
        writer = csv.writer(csvfile)
        writer.writerow(['Number', 'Text'])  # Writing header

        for number, text in matches:
            writer.writerow([number, text.strip()])  # Writing each row with number and text

    print(f"Data has been extracted and saved to {output_file}")


def extract_dates(input_file, output_file):
    extracted_data = []

    with open(input_file, 'r', encoding='utf-8') as csvfile:
        reader = csv.DictReader(csvfile)
        
        for row in reader:
            paragraph_number = row['Number']  # Get the paragraph number
            text = row['Text']  # Get the text
            
            # Use datefinder to find all dates in the text
            matches = datefinder.find_dates(text)
            
            # For each found date, add it to the extracted data list
            for match in matches:
                extracted_data.append({'Date': match, 'Text': text, 'Paragraph Number': paragraph_number})

    extracted_data.sort(key=lambda x: x['Date'])

    with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['Date', 'Text', 'Paragraph Number']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        
        writer.writeheader()
        
        for entry in extracted_data:
            writer.writerow({'Date': entry['Date'].strftime("%Y-%m-%d"), 
                             'Text': entry['Text'], 
                             'Paragraph Number': entry['Paragraph Number']})

    print(f"Date extraction completed and saved to {output_file}")

def create_word_document_from_csv(input_file, output_file):
    data = pd.read_csv(input_file)

    # Create a new Word document
    doc = Document()

    # Add a title to the document (optional)
    doc.add_heading('Draft Chronology', level=1)

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'  # Optional: set a table style

    # Add the header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Text'
    hdr_cells[2].text = 'Paragraph Number'

    for index, row in data.iterrows():
        row_cells = table.add_row().cells  # Add a new row
        row_cells[0].text = str(row['Date'])  # Fill Date
        row_cells[1].text = str(row['Text'])  # Fill Text
        row_cells[2].text = str(row['Paragraph Number'])  # Fill Paragraph Number

    doc.save(output_file)

    print(f"Word document '{output_file}' created successfully.")

def everything_function(f):
    clean_markdown_document(f, 'cleaned.md')
    extract_numbered_items_to_csv('cleaned.md', 'all_dates.csv')
    extract_dates('all_dates.csv', 'dates_extracted.csv')
    create_word_document_from_csv('dates_extracted.csv', 'output.docx')

    # Remove intermediate files
    for file in ['cleaned.md', 'all_dates.csv', 'dates_extracted.csv']:
        if os.path.exists(file):
            os.remove(file)
            print(f"Removed {file}")

@app.route('/')   
def main():
    everything_function() 

if __name__ == '__main__':   
    app.run(debug=True)